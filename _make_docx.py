# -*- coding: utf-8 -*-
"""거짓말 유형 도감을 워드 파일로 만든다."""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = '맑은 고딕'

FAMILIES = [
    {
        'id': 'A', 'name': '세기', 'slot': '얼마나',
        'sub': '얼마나 확실한가 · 얼마나 많은가',
        'note': '내용은 그대로 두고 문장의 온도만 올린다.',
        'rgb': (0x85, 0x4F, 0x0B), 'bg': 'FAEEDA',
        'types': [
            {'no': '①', 'name': '확신 올리기',
             'def': '저자가 "그럴 수 있다"고 조심스럽게 말한 것을 "그렇다"고 단정해 버린다.',
             'check': '원문에 「~일 수 있다 / ~로 보인다 / ~하는 경향이 있다 / 아마」가 있었는지 찾는다. 요약에서 그 꼬리가 잘렸으면 이 유형.',
             'orig': '결핍감이 4번 유형의 행동을 이끄는 요인일 수 있다.',
             'lie': '결핍감이 4번 유형의 행동을 이끄는 요인이다.',
             'vs': '③ 단서 떼기와 헷갈리면 — 사라진 말이 "확신을 낮추는 말"이면 ①, "상황을 한정하는 말"이면 ③.'},
            {'no': '②', 'name': '수량 부풀리기',
             'def': '몇몇을 대부분으로, 가끔을 자주로. 숫자 감각만 슬쩍 키운다.',
             'check': '원문의 수량어(일부·여럿·절반쯤·때때로)와 요약의 수량어(대체로·흔히·자주·거의)를 나란히 놓고 세어 본다.',
             'orig': '4번 유형 중 일부는 남의 관심을 부담스러워한다.',
             'lie': '4번 유형은 대체로 남의 관심을 부담스러워한다.',
             'vs': '① 확신 올리기와 헷갈리면 — "얼마나 확실한가"가 바뀌면 ①, "얼마나 많은가"가 바뀌면 ②.'},
            {'no': '③', 'name': '단서 떼기',
             'def': '원문에 붙어 있던 조건(~할 때만)이나 예외(다만 ~는 아니다)를 떼어내고 무조건인 것처럼 만든다.',
             'check': '원문에서 「만약 / ~하는 경우 / ~할 때 / 다만 / 하지만 / 물론 ~도 있다」를 찾는다. 그 대목이 요약에서 너무 깔끔하면 이 유형.',
             'orig': '결핍감이 커질 때 4번 유형은 자기 세계로 물러난다.',
             'lie': '4번 유형은 자기 세계로 물러난다.',
             'vs': '⑦ 대상 넓히기와 헷갈리면 — 사라진 것이 "언제·어떤 상황"이면 ③, "누구"면 ⑦.'},
        ]
    },
    {
        'id': 'B', 'name': '이음새', 'slot': '왜',
        'sub': '무엇이 무엇을 낳는가',
        'note': '문장 사이의 화살표를 건드린다. 단어는 거의 그대로라 눈으로는 안 잡힌다.',
        'rgb': (0x0F, 0x6E, 0x56), 'bg': 'E1F5EE',
        'types': [
            {'no': '④', 'name': '동행을 원인으로',
             'def': '"함께 나타난다"는 말을 "그래서 생긴다"는 말로 바꾼다.',
             'check': '원문이 「관련이 있다 / 함께 나타난다 / 동반된다」인데 요약이 「때문이다 / 만든다 / 낳는다」면 이 유형.',
             'orig': '결핍감이 강한 사람에게서 예술적 표현 욕구가 함께 관찰된다.',
             'lie': '결핍감이 예술적 표현 욕구를 만들어 낸다.',
             'vs': '⑤ 화살표 뒤집기와 헷갈리면 — 원문에 화살표가 아예 없었으면 ④, 있었는데 방향만 반대면 ⑤.'},
            {'no': '⑤', 'name': '화살표 뒤집기',
             'def': 'A 때문에 B라고 한 것을 B 때문에 A로 뒤집는다. 쓰인 단어는 하나도 바뀌지 않는다.',
             'check': '「때문에 / 그래서 / ~로 이어진다」 앞뒤에 무엇이 놓였는지 원문과 요약을 나란히 두고 화살표를 직접 그려 본다.',
             'orig': '채워지지 않은 결핍감 때문에 특별함에 집착하게 된다.',
             'lie': '특별함에 대한 집착 때문에 결핍감이 생긴다.',
             'vs': '단어가 전부 그대로라 읽어서는 못 잡는다. 반드시 화살표를 그려 방향을 맞춰 봐야 한다.'},
            {'no': '⑥', 'name': '이유 바꿔치기',
             'def': '결론은 원문 그대로 두고 "왜냐하면" 뒤만 다른 이유로 갈아 끼운다.',
             'check': '결론이 맞아 보이면 그냥 넘어가기 쉽다. 결론에 동그라미를 치고, 그 앞의 이유만 따로 떼어 원문과 맞춘다.',
             'orig': '4번 유형이 평범함을 견디지 못하는 것은 자기 정체성이 흔들리기 때문이다.',
             'lie': '4번 유형이 평범함을 견디지 못하는 것은 남보다 앞서고 싶기 때문이다.',
             'vs': '⑩ 결론 얹기와 헷갈리면 — 결론이 원문에 있던 것이면 ⑥, 원문에 아예 없던 것이면 ⑩.'},
        ]
    },
    {
        'id': 'C', 'name': '범위', 'slot': '어디까지·누가',
        'sub': '누구에게 · 어디까지 해당하는가',
        'note': '주장은 그대로 두고 울타리만 옮긴다.',
        'rgb': (0x3C, 0x34, 0x89), 'bg': 'EEEDFE',
        'types': [
            {'no': '⑦', 'name': '대상 넓히기',
             'def': '특정한 사람·상황에서 나온 이야기를 모두의 이야기로 넓힌다.',
             'check': '원문의 주인공이 정확히 누구였는지 짚는다. "4번 유형"이 "사람들"로, "어린 시절"이 "평생"으로 커졌으면 이 유형.',
             'orig': '4번 유형은 남과 다르다는 감각에서 자기를 확인한다.',
             'lie': '사람은 남과 다르다는 감각에서 자기를 확인한다.',
             'vs': '⑧ 주어 바꾸기와 헷갈리면 — 같은 대상이 커진 것이면 ⑦, 아예 다른 대상으로 옮겨 갔으면 ⑧.'},
            {'no': '⑧', 'name': '주어 바꾸기',
             'def': '하는 사람과 겪는 사람을 슬쩍 맞바꾼다.',
             'check': '그 행동을 "누가" 하는지 묻는다. 원문에서 아이가 한 일을 요약에서 부모가 하고 있지 않은지 확인.',
             'orig': '4번 유형은 상대가 자기를 알아봐 주기를 기다린다.',
             'lie': '4번 유형은 상대를 먼저 알아봐 주려고 한다.',
             'vs': '⑤ 화살표 뒤집기와 헷갈리면 — 바뀐 것이 "행동의 주인"이면 ⑧, "원인과 결과의 순서"면 ⑤.'},
            {'no': '⑨', 'name': '비교 기준 지우기',
             'def': '"무엇보다 낫다"에서 비교 대상을 지워, 그냥 좋은 것으로 만든다.',
             'check': '요약에 「더 / 낫다 / 강하다 / 높다 / 세밀하다」가 있으면 반드시 묻는다 — 무엇과 비교해서?',
             'orig': '4번 유형은 3번 유형보다 감정의 결을 세밀하게 느낀다.',
             'lie': '4번 유형은 감정의 결을 세밀하게 느끼는 유형이다.',
             'vs': '③ 단서 떼기와 헷갈리면 — 사라진 것이 "비교 대상"이면 ⑨, "조건·예외"면 ③.'},
        ]
    },
    {
        'id': 'D', 'name': '없던 말', 'slot': '없던 말',
        'sub': '원문에 없는 것을 끼워 넣는가',
        'note': '앞은 전부 진짜다. 마지막 한 조각만 저자가 하지 않은 말이다.',
        'rgb': (0x99, 0x3C, 0x1D), 'bg': 'FAECE7',
        'types': [
            {'no': '⑩', 'name': '결론 얹기',
             'def': '저자가 내리지 않은 판단이나 처방을 저자 말인 것처럼 문장 끝에 붙인다.',
             'check': '그 결론을 원문에서 찾아본다. 비슷한 문장조차 없으면 이 유형. 특히 「그러므로 ~해야 한다 / 결국 ~이다」를 의심한다.',
             'orig': '4번 유형은 결핍감 때문에 자기 세계로 물러나곤 한다. (원문은 여기서 끝난다)',
             'lie': '4번 유형은 결핍감 때문에 자기 세계로 물러나므로, 결핍감을 극복해야 성숙해진다.',
             'vs': '⑥ 이유 바꿔치기와 헷갈리면 — 원문에 있던 결론의 이유만 바뀌었으면 ⑥, 결론 자체가 새로 생겼으면 ⑩.'},
        ]
    },
]


def set_kfont(run, size=10.5, bold=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)


def shade(cell, hexcolor):
    el = OxmlElement('w:shd')
    el.set(qn('w:val'), 'clear')
    el.set(qn('w:fill'), hexcolor)
    cell._tc.get_or_add_tcPr().append(el)


def para(doc, text='', size=10.5, bold=False, color=None, space_after=6,
         align=None, indent=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.4
    if align:
        p.alignment = align
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    if text:
        set_kfont(p.add_run(text), size, bold, color)
    return p


def build(path):
    doc = Document()
    sec = doc.sections[0]
    sec.left_margin = sec.right_margin = Cm(2.2)
    sec.top_margin = sec.bottom_margin = Cm(2.0)

    st = doc.styles['Normal']
    st.font.name = FONT
    st.font.size = Pt(10.5)
    st.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)

    # ── 표지 ──
    para(doc, '거짓말 10종 도감', 22, True, space_after=2)
    para(doc, '책 분석 + 거짓말 사냥 · 비판적 독해 훈련 자료', 11, False, (0x6B, 0x6A, 0x65), 18)

    para(doc, '먼저 읽을 것', 14, True, (0x1E, 0x1E, 0x1B), 6)
    para(doc, '유형을 먼저 고르지 마세요. 요약 문장은 네 칸으로 이루어져 있고, '
              '거짓말은 그중 한 칸만 바꾼 것입니다. 그래서 바뀐 칸부터 정하고, '
              '그 칸에 해당하는 군 안에서만 유형을 고릅니다. '
              '10개 중에서 고르던 문제가 3개 중에서 고르는 문제로 줄어듭니다.', 10.5, space_after=12)

    para(doc, '찾는 순서', 12, True, space_after=4)
    for i, t in enumerate([
        '요약 문장 하나를 고르고, 본문에서 같은 내용이 나온 자리를 찾아 나란히 둔다.',
        '두 문장에서 바뀐 칸이 어디인지 정한다. (얼마나 / 왜 / 어디까지·누가 / 없던 말)',
        '그 칸에 해당하는 군 안에서만 유형을 고른다.',
    ], 1):
        para(doc, f'{i}. {t}', 10.5, space_after=3, indent=0.5)

    para(doc, '', space_after=8)

    # ── 한눈에 보기 표 ──
    para(doc, '한눈에 보기', 14, True, space_after=6)
    tb = doc.add_table(rows=1, cols=4)
    tb.style = 'Table Grid'
    tb.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = ['군', '바뀌는 칸', '무엇을 건드리나', '유형']
    for c, h in enumerate(hdr):
        cell = tb.rows[0].cells[c]
        cell.text = ''
        set_kfont(cell.paragraphs[0].add_run(h), 9.5, True, (0xFF, 0xFF, 0xFF))
        shade(cell, '2E2E2B')
    for f in FAMILIES:
        row = tb.add_row().cells
        vals = [f"{f['id']}군 · {f['name']}", f['slot'], f['note'],
                '  '.join(f"{t['no']} {t['name']}" for t in f['types'])]
        for c, v in enumerate(vals):
            row[c].text = ''
            set_kfont(row[c].paragraphs[0].add_run(v), 9, c == 0,
                      f['rgb'] if c == 0 else None)
        shade(row[0], f['bg'])

    for w, col in zip([Cm(2.3), Cm(2.6), Cm(6.0), Cm(6.1)], tb.columns):
        for cell in col.cells:
            cell.width = w

    doc.add_page_break()

    # ── 군별 상세 ──
    for f in FAMILIES:
        para(doc, f"{f['id']}군 · {f['name']}", 17, True, f['rgb'], 2)
        para(doc, f"바뀌는 칸 — {f['slot']}  |  {f['sub']}", 10, True, (0x6B, 0x6A, 0x65), 2)
        para(doc, f['note'], 10, False, (0x6B, 0x6A, 0x65), 10)

        for t in f['types']:
            para(doc, f"{t['no']} {t['name']}", 13, True, f['rgb'], 4)
            para(doc, t['def'], 10.5, space_after=6, indent=0.4)

            para(doc, '찾는 법', 9.5, True, (0x6B, 0x6A, 0x65), 2, indent=0.4)
            para(doc, t['check'], 10, space_after=8, indent=0.4)

            # 원문 / 비틀림 대조표
            cmp_tb = doc.add_table(rows=2, cols=2)
            cmp_tb.style = 'Table Grid'
            for r, (tag, txt, tagcolor, bgc) in enumerate([
                ('원문', t['orig'], (0x3B, 0x6D, 0x11), 'EAF3DE'),
                ('비틀림', t['lie'], (0xA3, 0x2D, 0x2D), 'FCEBEB'),
            ]):
                c0, c1 = cmp_tb.rows[r].cells
                c0.text = ''
                set_kfont(c0.paragraphs[0].add_run(tag), 9, True, tagcolor)
                shade(c0, bgc)
                c0.width = Cm(1.8)
                c1.text = ''
                set_kfont(c1.paragraphs[0].add_run(txt), 10)
                c1.width = Cm(15.2)

            para(doc, '', space_after=4)
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(14)
            p.paragraph_format.left_indent = Cm(0.4)
            p.paragraph_format.line_spacing = 1.4
            set_kfont(p.add_run('헷갈리지 않기  '), 9.5, True, (0x6B, 0x6A, 0x65))
            set_kfont(p.add_run(t['vs']), 10)

        if f['id'] != 'D':
            doc.add_page_break()

    # ── 출제 규칙 ──
    doc.add_page_break()
    para(doc, '거짓말을 만들 때의 규칙', 17, True, space_after=4)
    para(doc, 'AI가 거짓말 문장을 만들 때 지키는 규칙입니다. 사람이 직접 문제를 낼 때도 같습니다.',
         10, False, (0x6B, 0x6A, 0x65), 10)
    rules = [
        '거짓말은 정확히 3개, 서로 다른 군에서 하나씩. 같은 군 중복 금지.',
        '한 문장에서 바꾸는 건 딱 한 칸. 문장 전체를 반대로 만들지 않는다.',
        '바꾸지 않은 부분은 본문에 나온 표현 그대로. 새 단어·새 개념 등장 금지.',
        '극단 단어(모든/항상/절대/전혀) 금지. ② 수량 부풀리기도 "대체로/자주/흔히"까지만.',
        '거짓 문장은 참 문장과 길이·문체·어미가 구별되지 않아야 한다.',
        '참 문장 최소 2개에 "~일 수 있다", "~할 때", "~보다"를 미끼로 남긴다. '
        '이게 없으면 학생이 "조심스러운 표현이 없는 문장이 거짓말"이라는 요령만으로 다 찍는다.',
        '거짓말의 원래 내용은 본문 안에 실제로 있어야 한다. 없는 내용을 지어내 비틀지 않는다.',
        '힌트는 정답을 알려주지 않는다. 1단계는 무엇을 대조할지, 2단계는 위치 범위, 3단계는 스스로 던질 질문.',
    ]
    for i, r in enumerate(rules, 1):
        para(doc, f'{i}. {r}', 10.5, space_after=5, indent=0.4)

    para(doc, '', space_after=10)
    para(doc, '왜 이렇게 바꿨나', 14, True, space_after=6)
    para(doc, '이전 10종은 이름만 나열되어 있어 실제로 쓰기 어려웠습니다. '
              '인과 바꿔치기 / 상관-인과 둔갑 / 근거 바꿔치기는 학생 눈에 모두 "이유가 이상하다"로 읽혀 '
              '고를 수가 없었고, 조건 제거와 반례 무시, 범위 이동과 비교 대상 누락도 같은 문제를 가지고 있었습니다. '
              '무엇보다 AI에게는 유형 이름만 전달되고 정의가 전달되지 않아, 라벨과 맞지 않는 거짓말이 만들어졌습니다. '
              '지금은 화면에 보이는 정의와 AI가 받는 정의가 같은 하나의 원본에서 나옵니다.',
         10.5, space_after=8)

    doc.save(path)
    return path


if __name__ == '__main__':
    out = os.path.join(os.path.dirname(os.path.abspath(__file__)), '거짓말 10종 도감.docx')
    build(out)
    print('저장:', out, os.path.getsize(out), 'bytes')
