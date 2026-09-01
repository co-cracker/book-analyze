# -*- coding: utf-8 -*-
"""lie_types.json에서 인쇄용 워드 파일을 만든다.

    python make_docx.py

유형을 고칠 때는 lie_types.json만 고치고 build.py와 이걸 차례로 돌리면
앱·웹·워드가 모두 같은 내용이 된다.
"""
import json
import os

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
FONT = '맑은 고딕'


def rgb(hexcolor):
    return tuple(int(hexcolor[i:i + 2], 16) for i in (1, 3, 5))


def set_kfont(run, size=10.5, bold=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)


def shade(cell, hexfill):
    el = OxmlElement('w:shd')
    el.set(qn('w:val'), 'clear')
    el.set(qn('w:fill'), hexfill)
    cell._tc.get_or_add_tcPr().append(el)


def para(doc, text='', size=10.5, bold=False, color=None,
         space_after=6, indent=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.4
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    if text:
        set_kfont(p.add_run(text), size, bold, color)
    return p


def labelled(doc, label, text, indent=0.4, space_after=6):
    """작은 라벨 + 본문을 한 단락에 넣는다."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.4
    p.paragraph_format.left_indent = Cm(indent)
    set_kfont(p.add_run(label + '  '), 9, True, (0x7A, 0x76, 0x6C))
    set_kfont(p.add_run(text), 10)
    return p


GREY = (0x6B, 0x6A, 0x65)
INK = (0x1B, 0x1A, 0x17)


def build(data, path):
    doc = Document()
    sec = doc.sections[0]
    sec.left_margin = sec.right_margin = Cm(2.2)
    sec.top_margin = sec.bottom_margin = Cm(2.0)

    st = doc.styles['Normal']
    st.font.name = FONT
    st.font.size = Pt(10.5)
    st.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)

    fams = data['families']
    total = sum(len(f['types']) for f in fams)

    # ── 표지 ──
    para(doc, data['title'], 22, True, space_after=2)
    para(doc, data['subtitle'], 11, False, GREY, 16)

    o = data['origin']
    para(doc, o['title'], 14, True, space_after=4)
    para(doc, o['lead'], 12, True, INK, 4)
    para(doc, ' → '.join(o['flow']) + '   (' + o['flowNote'] + ')', 10, False, GREY, 4)
    para(doc, o['body'], 10.5, False, None, 14)

    para(doc, '유형을 고르기 전에 — 네 개의 질문', 14, True, space_after=4)
    para(doc, data['decision']['question'] + ' ' + data['decision']['note'], 10.5, False, GREY, 10)

    # ── 4군 판별표 ──
    tb = doc.add_table(rows=1, cols=3)
    tb.style = 'Table Grid'
    tb.alignment = WD_TABLE_ALIGNMENT.CENTER
    for c, h in enumerate(['군', '던지는 질문', '이럴 때 걸린다']):
        cell = tb.rows[0].cells[c]
        cell.text = ''
        set_kfont(cell.paragraphs[0].add_run(h), 9.5, True, (0xFF, 0xFF, 0xFF))
        shade(cell, '2E2E2B')
    for f in fams:
        row = tb.add_row().cells
        for c, v in enumerate(['%s군 · %s' % (f['id'], f['name']), f['ask'], f['trigger']]):
            row[c].text = ''
            set_kfont(row[c].paragraphs[0].add_run(v), 9, c == 0,
                      rgb(f['color']) if c == 0 else None)
        shade(row[0], f['bg'].lstrip('#'))
    for w, col in zip([Cm(3.2), Cm(6.2), Cm(7.6)], tb.columns):
        for cell in col.cells:
            cell.width = w

    para(doc, '', space_after=4)
    para(doc, '각 군에 유형이 3개씩 있습니다. 위 표에서 군을 먼저 정하면 %d개 중에서 고르던 문제가 '
              '3개 중에서 고르는 문제로 줄어듭니다.' % total, 10, False, GREY, 10)

    th = data['threshold']
    para(doc, th['title'], 13, True, space_after=4)
    para(doc, th['lead'], 10.5, False, None, 6)
    para(doc, th['rule'], 12, True, (0xB5, 0x37, 0x2B), 2, indent=0.4)
    para(doc, th['sub'], 10, False, GREY, 6, indent=0.4)

    # ── 군별 상세 ──
    for f in fams:
        doc.add_page_break()
        color = rgb(f['color'])
        para(doc, '%s군 · %s' % (f['id'], f['name']), 17, True, color, 2)
        para(doc, f['ask'], 12, True, color, 2)
        para(doc, '이럴 때 걸린다 — ' + f['trigger'], 10, False, GREY, 4)
        if f.get('note'):
            para(doc, f['note'], 10, False, GREY, 12, indent=0.4)
        else:
            para(doc, '', space_after=6)

        for t in f['types']:
            para(doc, '%s %s' % (t['no'], t['name']), 13.5, True, color, 3)
            labelled(doc, '정의', t['def'])
            labelled(doc, '알아보는 법', t['spot'], space_after=8)

            def pair_table(rows, w0=Cm(2.7), w1=Cm(14.3)):
                tb2 = doc.add_table(rows=len(rows), cols=2)
                tb2.style = 'Table Grid'
                for r, (tag, txt, tagcolor, fill, bold) in enumerate(rows):
                    c0, c1 = tb2.rows[r].cells
                    c0.text = ''
                    set_kfont(c0.paragraphs[0].add_run(tag), 8.5, True, tagcolor)
                    shade(c0, fill)
                    c0.width = w0
                    c1.text = ''
                    set_kfont(c1.paragraphs[0].add_run(txt), 10, bold)
                    c1.width = w1

            pairs = [('작가', t['bookAuthor'], (0x2F, 0x6B, 0x4F), 'EAF3EC', False),
                     ('정리', t['bookSummary'], (0xB5, 0x37, 0x2B), 'FBEDEA', True)]
            if t.get('bookAuthor2'):
                pairs += [('작가', t['bookAuthor2'], (0x2F, 0x6B, 0x4F), 'EAF3EC', False),
                          (t.get('summaryLabel2', '정리'), t['bookSummary2'],
                           (0xB5, 0x37, 0x2B), 'FBEDEA', True)]
            pair_table(pairs)
            para(doc, '', space_after=6)

            if t.get('variant'):
                labelled(doc, '변형', t['variant'])

            labelled(doc, '거짓말이 되는 선', t['line'])
            labelled(doc, '믿으면 벌어지는 일', t['believe'])

            para(doc, '글 읽을 때 만나는 자리', 9, True, GREY, 3, indent=0.4)
            for pl in t['places']:
                para(doc, '· ' + pl, 10, False, None, 3, indent=0.7)
            para(doc, '', space_after=4)

            q = doc.add_paragraph()
            q.paragraph_format.space_after = Pt(8)
            q.paragraph_format.left_indent = Cm(0.4)
            q.paragraph_format.line_spacing = 1.4
            set_kfont(q.add_run('되받아치는 질문   '), 8.5, True, GREY)
            set_kfont(q.add_run('“' + t['ask'] + '”'), 11.5, True, color)

            if t.get('vs'):
                labelled(doc, '헷갈리지 않기', t['vs'], space_after=16)
            else:
                para(doc, '', space_after=10)

    # ── 규칙 ──
    doc.add_page_break()
    ru = data['rules']
    para(doc, '문제를 만드는 규칙', 17, True, space_after=6)

    para(doc, ru['passTitle'] + ' — ' + ru['passNote'], 12, True, None, 5)
    for a, b in ru['pass']:
        labelled(doc, a, b)
    para(doc, '', space_after=6)

    para(doc, ru['makeTitle'], 12, True, None, 5)
    for i, r in enumerate(ru['make'], 1):
        para(doc, '%d. %s' % (i, r), 10.5, space_after=5, indent=0.4)
    para(doc, '', space_after=6)

    para(doc, ru['hintTitle'], 12, True, None, 5)
    for a, b in ru['hints']:
        labelled(doc, a, b)
    para(doc, '', space_after=6)

    para(doc, ru['cardTitle'], 12, True, None, 5)
    for a, b in ru['card']:
        labelled(doc, a, b)

    cl = data['closing']
    para(doc, '', space_after=10)
    para(doc, cl['title'], 14, True, space_after=4)
    para(doc, cl['lead'], 10.5, False, GREY, 4)
    para(doc, '“' + cl['question'] + '”', 13, True, INK, 6, indent=0.4)
    para(doc, cl['body'], 10.5, False, None, 6)

    doc.save(path)
    return path


def main():
    with open(os.path.join(HERE, 'lie_types.json'), encoding='utf-8') as f:
        data = json.load(f)
    out = os.path.join(HERE, data['title'] + '.docx')
    build(data, out)
    print('저장:', os.path.basename(out), os.path.getsize(out), 'bytes')


if __name__ == '__main__':
    main()
